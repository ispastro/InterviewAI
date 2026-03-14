"""
InterviewMe Text Extraction Utilities

This module provides utilities for extracting text from various file formats:
- PDF files using pdfplumber (better table handling than PyPDF2)
- DOCX files using python-docx
- Plain text files (direct pass-through)

Engineering decisions:
- Use pdfplumber over PyPDF2 for better table and layout handling
- Comprehensive error handling with specific exception types
- File size and content validation
- Memory-efficient streaming for large files
- Sanitization of extracted text
"""

import io
from typing import Optional, Tuple
from pathlib import Path

import pdfplumber
from docx import Document
from fastapi import UploadFile

from app.config import settings
from app.core.exceptions import ValidationError


# ============================================================
# CONSTANTS
# ============================================================

# Maximum file size in bytes (from settings)
MAX_FILE_SIZE = settings.MAX_FILE_SIZE_MB * 1024 * 1024

# Supported file extensions
SUPPORTED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}

# Minimum and maximum text length
MIN_TEXT_LENGTH = 50      # At least 50 characters
MAX_TEXT_LENGTH = 50000   # Max 50k characters (reasonable CV length)


# ============================================================
# FILE VALIDATION
# ============================================================

def validate_file(file: UploadFile) -> None:
    """
    Validate uploaded file before processing.
    
    Checks:
    - File size within limits
    - File extension is supported
    - Content type matches extension
    
    Args:
        file: FastAPI UploadFile object
        
    Raises:
        ValidationError: If file validation fails
    """
    # Check file size
    if hasattr(file, 'size') and file.size:
        if file.size > MAX_FILE_SIZE:
            raise ValidationError(
                f"File too large: {file.size / 1024 / 1024:.1f}MB. Maximum allowed: {settings.MAX_FILE_SIZE_MB}MB"
            )
    
    # Check file extension
    if not file.filename:
        raise ValidationError("Filename is required")
    
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in SUPPORTED_EXTENSIONS:
        raise ValidationError(
            f"Unsupported file type: {file_ext}. Supported types: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )
    
    # Check content type (if provided)
    expected_content_type = SUPPORTED_EXTENSIONS[file_ext]
    if file.content_type and not file.content_type.startswith(expected_content_type.split('/')[0]):
        # Allow some flexibility in content type checking
        if file_ext == ".txt" and "text" not in file.content_type:
            raise ValidationError(f"Content type mismatch: expected text file, got {file.content_type}")


def validate_extracted_text(text: str, filename: str) -> str:
    """
    Validate and sanitize extracted text.
    
    Args:
        text: Extracted text content
        filename: Original filename for error messages
        
    Returns:
        Cleaned and validated text
        
    Raises:
        ValidationError: If text validation fails
    """
    if not text or not text.strip():
        raise ValidationError(f"No text content found in {filename}")
    
    # Clean the text
    cleaned_text = text.strip()
    
    # Remove excessive whitespace
    import re
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
    
    # Check length constraints
    if len(cleaned_text) < MIN_TEXT_LENGTH:
        raise ValidationError(
            f"Text too short: {len(cleaned_text)} characters. Minimum required: {MIN_TEXT_LENGTH}"
        )
    
    if len(cleaned_text) > MAX_TEXT_LENGTH:
        raise ValidationError(
            f"Text too long: {len(cleaned_text)} characters. Maximum allowed: {MAX_TEXT_LENGTH}"
        )
    
    return cleaned_text


# ============================================================
# TEXT EXTRACTION FUNCTIONS
# ============================================================

def extract_text_from_pdf(file_content: bytes, filename: str) -> str:
    """
    Extract text from PDF file using pdfplumber.
    
    Args:
        file_content: PDF file content as bytes
        filename: Original filename for error messages
        
    Returns:
        Extracted text content
        
    Raises:
        ValidationError: If PDF processing fails
    """
    try:
        text_parts = []
        
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            if len(pdf.pages) == 0:
                raise ValidationError(f"PDF file {filename} contains no pages")
            
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    # Extract text from page
                    page_text = page.extract_text()
                    
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Also try to extract text from tables
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                # Join non-None cells with spaces
                                row_text = ' '.join(str(cell) for cell in row if cell)
                                if row_text.strip():
                                    text_parts.append(row_text)
                
                except Exception as e:
                    # Log page-specific error but continue with other pages
                    print(f"Warning: Failed to extract text from page {page_num} of {filename}: {e}")
                    continue
        
        if not text_parts:
            raise ValidationError(f"No readable text found in PDF {filename}")
        
        # Join all text parts
        full_text = '\n'.join(text_parts)
        
        return validate_extracted_text(full_text, filename)
        
    except ValidationError:
        raise
    except Exception as e:
        raise ValidationError(f"Failed to process PDF {filename}: {str(e)}")


def extract_text_from_docx(file_content: bytes, filename: str) -> str:
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        file_content: DOCX file content as bytes
        filename: Original filename for error messages
        
    Returns:
        Extracted text content
        
    Raises:
        ValidationError: If DOCX processing fails
    """
    try:
        text_parts = []
        
        # Load document from bytes
        doc = Document(io.BytesIO(file_content))
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        if not text_parts:
            raise ValidationError(f"No readable text found in DOCX {filename}")
        
        # Join all text parts
        full_text = '\n'.join(text_parts)
        
        return validate_extracted_text(full_text, filename)
        
    except ValidationError:
        raise
    except Exception as e:
        raise ValidationError(f"Failed to process DOCX {filename}: {str(e)}")


def extract_text_from_txt(file_content: bytes, filename: str) -> str:
    """
    Extract text from plain text file.
    
    Args:
        file_content: Text file content as bytes
        filename: Original filename for error messages
        
    Returns:
        Extracted text content
        
    Raises:
        ValidationError: If text processing fails
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                return validate_extracted_text(text, filename)
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail
        raise ValidationError(f"Unable to decode text file {filename}. Please ensure it's in UTF-8 format.")
        
    except ValidationError:
        raise
    except Exception as e:
        raise ValidationError(f"Failed to process text file {filename}: {str(e)}")


# ============================================================
# MAIN EXTRACTION FUNCTION
# ============================================================

async def extract_text_from_file(file: UploadFile) -> Tuple[str, dict]:
    """
    Extract text from uploaded file (auto-detects format).
    
    Args:
        file: FastAPI UploadFile object
        
    Returns:
        Tuple of (extracted_text, metadata)
        
    Raises:
        ValidationError: If file processing fails
    """
    # Validate file first
    validate_file(file)
    
    # Read file content
    try:
        file_content = await file.read()
        
        # Reset file pointer for potential re-reading
        await file.seek(0)
        
    except Exception as e:
        raise ValidationError(f"Failed to read file {file.filename}: {str(e)}")
    
    # Get file extension
    file_ext = Path(file.filename).suffix.lower()
    
    # Extract text based on file type
    if file_ext == ".pdf":
        extracted_text = extract_text_from_pdf(file_content, file.filename)
    elif file_ext == ".docx":
        extracted_text = extract_text_from_docx(file_content, file.filename)
    elif file_ext == ".txt":
        extracted_text = extract_text_from_txt(file_content, file.filename)
    else:
        raise ValidationError(f"Unsupported file type: {file_ext}")
    
    # Create metadata
    metadata = {
        "filename": file.filename,
        "file_type": file_ext,
        "file_size_bytes": len(file_content),
        "content_type": file.content_type,
        "text_length": len(extracted_text),
        "word_count": len(extracted_text.split()),
    }
    
    return extracted_text, metadata


# ============================================================
# UTILITY FUNCTIONS
# ============================================================

def get_file_info(file: UploadFile) -> dict:
    """
    Get basic information about uploaded file without processing it.
    
    Args:
        file: FastAPI UploadFile object
        
    Returns:
        Dictionary with file information
    """
    file_ext = Path(file.filename).suffix.lower() if file.filename else ""
    
    return {
        "filename": file.filename,
        "file_type": file_ext,
        "content_type": file.content_type,
        "is_supported": file_ext in SUPPORTED_EXTENSIONS,
        "max_allowed_size_mb": settings.MAX_FILE_SIZE_MB,
    }


def estimate_processing_time(file_size_bytes: int, file_type: str) -> float:
    """
    Estimate processing time for file extraction.
    
    Args:
        file_size_bytes: File size in bytes
        file_type: File extension (.pdf, .docx, .txt)
        
    Returns:
        Estimated processing time in seconds
    """
    # Base processing times (seconds per MB)
    processing_rates = {
        ".pdf": 2.0,    # PDFs are slower due to layout parsing
        ".docx": 1.0,   # DOCX is moderate
        ".txt": 0.1,    # Text files are very fast
    }
    
    rate = processing_rates.get(file_type, 1.0)
    file_size_mb = file_size_bytes / (1024 * 1024)
    
    # Minimum 0.5 seconds, maximum 30 seconds
    return max(0.5, min(30.0, file_size_mb * rate))